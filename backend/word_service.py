from docx import Document
from docx.shared import Pt
from io import BytesIO

def generar_word_sesion(datos, contenido):
    doc = Document()

    doc.add_heading("Sesión de Aprendizaje", 0)

    doc.add_paragraph(f"Nivel: {datos['nivel']}")
    doc.add_paragraph(f"Grado: {datos['grado']}")
    doc.add_paragraph(f"Área: {datos['area']}")
    doc.add_paragraph(f"Tema: {datos['tema']}")

    doc.add_paragraph("")
    doc.add_heading("Contenido generado por IA", level=1)

    doc.add_paragraph(contenido)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
